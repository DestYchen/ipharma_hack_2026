from __future__ import annotations

import argparse
import json
import os
import re
import sqlite3
import uuid
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

import requests
from docx import Document


DB_PATH = Path(__file__).resolve().parent / "pharma_runs.db"
PROMPT_PATH_DEFAULT = Path(__file__).resolve().parent / "промт для синопсиса.txt"
TEMPLATE_PATH_DEFAULT = Path(__file__).resolve().parent / "Синопсис для гпт ver.2 (3).docx"

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = os.getenv("OPENROUTER_MODEL", "openai/gpt-5.2")


@dataclass
class RunRecord:
    id: str
    created_at: str
    status: str
    mode: str
    session_id: str | None
    query: dict[str, Any] | None
    selected_reference_drug: str | None
    selection_rows_count: int | None
    selection_payload: dict[str, Any] | None
    router_output_text: str | None


def _init_db() -> None:
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute("PRAGMA journal_mode=WAL;")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS synopsis_runs (
                id TEXT PRIMARY KEY,
                created_at TEXT NOT NULL,
                status TEXT NOT NULL,
                source_run_id TEXT NOT NULL,
                template_path TEXT,
                prompt_path TEXT,
                attributes_json TEXT,
                output_markdown TEXT,
                output_docx_path TEXT,
                error_text TEXT
            )
            """
        )
        conn.commit()
    finally:
        conn.close()


def _load_run(run_id: str) -> RunRecord:
    conn = sqlite3.connect(DB_PATH)
    try:
        cur = conn.execute(
            """
            SELECT id, created_at, status, mode, session_id, query_json,
                   selected_reference_drug, selection_rows_count, selection_json, router_output_text
            FROM runs
            WHERE id = ?
            """,
            (run_id,),
        )
        row = cur.fetchone()
    finally:
        conn.close()
    if not row:
        raise KeyError(f"Run not found: {run_id}")
    (
        rid,
        created_at,
        status,
        mode,
        session_id,
        query_json,
        selected_reference_drug,
        selection_rows_count,
        selection_json,
        router_output_text,
    ) = row
    query = json.loads(query_json) if query_json else None
    selection_payload = json.loads(selection_json) if selection_json else None
    return RunRecord(
        id=rid,
        created_at=created_at,
        status=status,
        mode=mode,
        session_id=session_id,
        query=query,
        selected_reference_drug=selected_reference_drug,
        selection_rows_count=selection_rows_count,
        selection_payload=selection_payload,
        router_output_text=router_output_text,
    )


def _store_synopsis_run(
    *,
    status: str,
    source_run_id: str,
    template_path: str | None,
    prompt_path: str | None,
    attributes: dict[str, Any] | None,
    output_markdown: str | None,
    output_docx_path: str | None,
    error_text: str | None,
) -> str:
    run_id = uuid.uuid4().hex
    created_at = datetime.now().isoformat(timespec="seconds")
    attrs_json = json.dumps(attributes, ensure_ascii=False) if attributes is not None else None
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            """
            INSERT INTO synopsis_runs (
                id, created_at, status, source_run_id, template_path, prompt_path,
                attributes_json, output_markdown, output_docx_path, error_text
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                run_id,
                created_at,
                status,
                source_run_id,
                template_path,
                prompt_path,
                attrs_json,
                output_markdown,
                output_docx_path,
                error_text,
            ),
        )
        conn.commit()
    finally:
        conn.close()
    return run_id


def _update_synopsis_run(run_id: str, **fields: Any) -> None:
    if not fields:
        return
    allowed = {
        "status",
        "attributes_json",
        "output_markdown",
        "output_docx_path",
        "error_text",
    }
    updates = {k: v for k, v in fields.items() if k in allowed}
    if not updates:
        return
    set_clause = ", ".join(f"{k} = ?" for k in updates.keys())
    values = list(updates.values()) + [run_id]
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(f"UPDATE synopsis_runs SET {set_clause} WHERE id = ?", values)
        conn.commit()
    finally:
        conn.close()


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _collect_attributes(run: RunRecord) -> dict[str, Any]:
    attrs: dict[str, Any] = {}

    if run.query:
        attrs.update({
            "mnn": run.query.get("mnn"),
            "routes": run.query.get("routes"),
            "base_form": run.query.get("base_form"),
            "release_type": run.query.get("release_type"),
            "dosage": run.query.get("dosage"),
        })

    if run.selected_reference_drug:
        attrs["reference_drug"] = run.selected_reference_drug

    if run.selection_payload:
        attrs["selection_payload"] = run.selection_payload

        rows = run.selection_payload.get("selected_reference_rows") if isinstance(run.selection_payload, dict) else None
        if isinstance(rows, list) and rows:
            # Example: take first row as representative.
            row0 = rows[0]
            for key in ("trade_name", "drug_form", "dosage", "country", "ru_number", "ru_date", "mnn"):
                if key in row0 and row0[key] is not None:
                    attrs.setdefault(key, row0[key])

    if run.router_output_text:
        attrs["router_output_text"] = run.router_output_text

    return attrs


def _build_prompt(prompt_text: str, attributes: dict[str, Any], template_text: str) -> str:
    return (
        prompt_text
        .replace("<<<ATTRIBUTES>>>", json.dumps(attributes, ensure_ascii=False))
        .replace("<<<SYNOPSIS_TEMPLATE>>>", template_text)
    )


def _openrouter_chat(prompt: str) -> str:
    api_key = "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"
    if not api_key:
        raise RuntimeError("Set OPENROUTER_API_KEY env var")

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": os.getenv("OPENROUTER_REFERER", "https://local-synopsis"),
        "X-Title": os.getenv("OPENROUTER_TITLE", "Pharma Synopsis Builder"),
    }
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": "Ты аккуратный медицинский аналитик. Не выдумывай факты."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 7000,
    }
    r = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=180)
    r.raise_for_status()
    data = r.json()
    return data["choices"][0]["message"]["content"]


def _find_markdown_table(text: str) -> list[list[str]] | None:
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    start = -1
    for i in range(len(lines) - 1):
        if "|" in lines[i] and re.match(r"^\s*\|?\s*-+", lines[i + 1]):
            start = i
            break
    if start == -1:
        return None

    def split_row(line: str) -> list[str]:
        row = line.strip().strip("|")
        return [c.strip() for c in row.split("|")]

    header = split_row(lines[start])
    rows: list[list[str]] = []
    for j in range(start + 2, len(lines)):
        if "|" not in lines[j]:
            break
        rows.append(split_row(lines[j]))
    if not header or not rows:
        return None
    return [header] + rows


def _write_docx_from_table(table_rows: list[list[str]], output_path: Path) -> None:
    doc = Document()
    doc.add_heading("Синопсис", level=1)

    cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"

    # Header (bold)
    hdr_cells = table.rows[0].cells
    for i, val in enumerate(table_rows[0]):
        run = hdr_cells[i].paragraphs[0].add_run(val)
        run.bold = True

    for row in table_rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(val)
            if i == 1:
                run.bold = True

    doc.save(str(output_path))


def build_synopsis(
    *,
    run_id: str,
    prompt_path: Path,
    template_path: Path,
    output_docx_path: Path,
    ) -> tuple[str, str, str]:
    _init_db()
    run = _load_run(run_id)

    prompt_text = _read_text(prompt_path)
    template_text = _extract_docx_text(template_path)
    attributes = _collect_attributes(run)

    synopsis_run_id = _store_synopsis_run(
        status="running",
        source_run_id=run_id,
        template_path=str(template_path),
        prompt_path=str(prompt_path),
        attributes=attributes,
        output_markdown=None,
        output_docx_path=None,
        error_text=None,
    )

    try:
        final_prompt = _build_prompt(prompt_text, attributes, template_text)
        markdown = _openrouter_chat(final_prompt)

        table = _find_markdown_table(markdown)
        if not table:
            raise RuntimeError("Model output does not contain a markdown table")

        _write_docx_from_table(table, output_docx_path)

        _update_synopsis_run(
            synopsis_run_id,
            status="done",
            output_markdown=markdown,
            output_docx_path=str(output_docx_path),
        )
        return markdown, str(output_docx_path), synopsis_run_id
    except Exception as exc:
        _update_synopsis_run(
            synopsis_run_id,
            status="error",
            error_text=str(exc),
        )
        raise


def build_synopsis_for_run(
    *,
    run_id: str,
    prompt_path: Path | None = None,
    template_path: Path | None = None,
    output_docx_path: Path | None = None,
) -> dict[str, Any]:
    prompt_path = prompt_path or PROMPT_PATH_DEFAULT
    template_path = template_path or TEMPLATE_PATH_DEFAULT
    if output_docx_path is None:
        output_docx_path = Path(f"synopsis_{run_id}.docx")
    markdown, docx_path, synopsis_run_id = build_synopsis(
        run_id=run_id,
        prompt_path=prompt_path,
        template_path=template_path,
        output_docx_path=output_docx_path,
    )
    return {
        "run_id": run_id,
        "synopsis_run_id": synopsis_run_id,
        "output_docx_path": docx_path,
        "output_markdown": markdown,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Build synopsis from DB + template")
    parser.add_argument("--run-id", required=True, help="Run id from pharma_runs.db")
    parser.add_argument("--prompt", default=str(PROMPT_PATH_DEFAULT), help="Path to prompt text")
    parser.add_argument("--template", default=str(TEMPLATE_PATH_DEFAULT), help="Path to synopsis template docx")
    parser.add_argument("--output-docx", default="synopsis_output.docx", help="Output docx file")
    args = parser.parse_args()

    markdown, out_path, _synopsis_run_id = build_synopsis(
        run_id=args.run_id,
        prompt_path=Path(args.prompt),
        template_path=Path(args.template),
        output_docx_path=Path(args.output_docx),
    )
    print("Saved:", out_path)
    print("\n--- MARKDOWN ---\n")
    print(markdown)


if __name__ == "__main__":
    main()
